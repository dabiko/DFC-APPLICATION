"""
Text extraction utilities for various document formats.

Supports extraction from:
- PDF documents (PyPDF2)
- Word documents (.docx) (python-docx)
- Excel spreadsheets (.xlsx) (openpyxl)
- Plain text files
- Images via OCR (pytesseract)

Also includes utilities for detecting scanned documents and
preprocessing images for better OCR accuracy.
"""

import PyPDF2
from docx import Document as DocxDocument
import openpyxl
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Extract text from various document formats.

    Provides static methods for extracting text from different file types,
    with comprehensive error handling and logging.
    """

    @staticmethod
    def extract_from_pdf(file_path):
        """
        Extract text from PDF document.

        Uses PyPDF2 to extract text from all pages of a PDF.
        For scanned PDFs with no extractable text, use OCR instead.

        Args:
            file_path (str): Path to PDF file

        Returns:
            str: Extracted text from all pages

        Raises:
            Exception: If PDF extraction fails
        """
        text = ""
        try:
            logger.info(f"Extracting text from PDF: {file_path}")

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                logger.info(f"PDF has {num_pages} pages")

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num} ---\n"
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue

                # Check if we got any text
                if not text.strip():
                    logger.warning(f"No text extracted from PDF (possibly scanned): {file_path}")
                    return ""

                logger.info(f"Successfully extracted {len(text)} characters from PDF")
                return text.strip()

        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}", exc_info=True)
            raise Exception(f"PDF extraction failed: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path):
        """
        Extract text from Word document (.docx).

        Extracts text from paragraphs, including tables.

        Args:
            file_path (str): Path to DOCX file

        Returns:
            str: Extracted text from document

        Raises:
            Exception: If DOCX extraction fails
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")

            doc = DocxDocument(file_path)

            # Extract text from paragraphs
            paragraphs_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs_text.append(paragraph.text)

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))

            # Combine all text
            all_text = []

            if paragraphs_text:
                all_text.append("--- Document Content ---")
                all_text.extend(paragraphs_text)

            if tables_text:
                all_text.append("\n--- Tables ---")
                all_text.extend(tables_text)

            text = "\n".join(all_text)

            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()

        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}", exc_info=True)
            raise Exception(f"DOCX extraction failed: {str(e)}")

    @staticmethod
    def extract_from_xlsx(file_path):
        """
        Extract text from Excel spreadsheet (.xlsx).

        Extracts values from all worksheets, preserving sheet structure.

        Args:
            file_path (str): Path to XLSX file

        Returns:
            str: Extracted text from all sheets

        Raises:
            Exception: If XLSX extraction fails
        """
        try:
            logger.info(f"Extracting text from XLSX: {file_path}")

            # Load workbook (data_only=True to get calculated values)
            wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)

            all_text = []

            for sheet in wb.worksheets:
                sheet_text = [f"\n--- Sheet: {sheet.title} ---"]

                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None and str(cell).strip()]

                    if row_values:
                        sheet_text.append(" | ".join(row_values))

                if len(sheet_text) > 1:  # More than just the header
                    all_text.extend(sheet_text)

            text = "\n".join(all_text)

            logger.info(f"Successfully extracted {len(text)} characters from XLSX ({len(wb.worksheets)} sheets)")
            wb.close()

            return text.strip()

        except Exception as e:
            logger.error(f"XLSX extraction failed for {file_path}: {e}", exc_info=True)
            raise Exception(f"XLSX extraction failed: {str(e)}")

    @staticmethod
    def extract_from_txt(file_path):
        """
        Extract text from plain text file.

        Supports various encodings with fallback.

        Args:
            file_path (str): Path to text file

        Returns:
            str: File contents

        Raises:
            Exception: If text file reading fails
        """
        try:
            logger.info(f"Extracting text from TXT: {file_path}")

            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Successfully read text file with {encoding} encoding")
                    return text.strip()
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

            logger.warning(f"Read text file with error handling (some characters may be lost)")
            return text.strip()

        except Exception as e:
            logger.error(f"Text file extraction failed for {file_path}: {e}", exc_info=True)
            raise Exception(f"Text file extraction failed: {str(e)}")

    @staticmethod
    def extract_text(file_path, file_type):
        """
        Extract text based on file type (dispatcher method).

        Routes to appropriate extraction method based on MIME type.

        Args:
            file_path (str): Path to file
            file_type (str): MIME type of file

        Returns:
            str: Extracted text, or empty string if unsupported format

        Example:
            >>> text = TextExtractor.extract_text('/path/to/file.pdf', 'application/pdf')
        """
        logger.info(f"Extracting text from {file_path} (type: {file_type})")

        # Map MIME types to extraction methods
        extractors = {
            'application/pdf': TextExtractor.extract_from_pdf,

            # Word documents
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': TextExtractor.extract_from_docx,
            'application/msword': TextExtractor.extract_from_docx,  # .doc (if supported)

            # Excel spreadsheets
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': TextExtractor.extract_from_xlsx,
            'application/vnd.ms-excel': TextExtractor.extract_from_xlsx,  # .xls (if supported)

            # Text files
            'text/plain': TextExtractor.extract_from_txt,
            'text/csv': TextExtractor.extract_from_txt,
            'application/csv': TextExtractor.extract_from_txt,
        }

        extractor = extractors.get(file_type)

        if not extractor:
            logger.warning(f"No extractor available for file type: {file_type}")
            return ""

        try:
            return extractor(file_path)
        except Exception as e:
            logger.error(f"Extraction failed: {e}", exc_info=True)
            # Return empty string instead of raising to allow processing to continue
            return ""

    @staticmethod
    def get_file_size_mb(file_path):
        """
        Get file size in megabytes.

        Args:
            file_path (str): Path to file

        Returns:
            float: File size in MB
        """
        try:
            size_bytes = os.path.getsize(file_path)
            return round(size_bytes / (1024 * 1024), 2)
        except Exception:
            return 0.0


class OCRDetector:
    """
    Utilities for detecting scanned documents that require OCR.
    """

    @staticmethod
    def is_scanned_document(file_path, file_type):
        """
        Detect if PDF is scanned (contains no selectable text).

        A document is considered scanned if:
        1. It's a PDF
        2. Text extraction yields very little text (< 50 characters)

        Args:
            file_path (str): Path to file
            file_type (str): MIME type

        Returns:
            bool: True if document appears to be scanned
        """
        # Only PDFs can be scanned
        if file_type != 'application/pdf':
            return False

        try:
            logger.info(f"Checking if PDF is scanned: {file_path}")

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Check first few pages for text
                pages_to_check = min(3, len(pdf_reader.pages))
                total_text = ""

                for i in range(pages_to_check):
                    try:
                        page_text = pdf_reader.pages[i].extract_text().strip()
                        total_text += page_text
                    except Exception:
                        continue

                # If very little text extracted, likely scanned
                is_scanned = len(total_text) < 50

                logger.info(
                    f"PDF {'appears to be scanned' if is_scanned else 'contains selectable text'} "
                    f"({len(total_text)} characters extracted from {pages_to_check} pages)"
                )

                return is_scanned

        except Exception as e:
            logger.error(f"Error detecting scanned document: {e}", exc_info=True)
            # Assume it's scanned if we can't determine
            return True


class ImagePreprocessor:
    """
    Image preprocessing utilities for improving OCR accuracy.
    """

    @staticmethod
    def preprocess_for_ocr(image):
        """
        Preprocess image for better OCR accuracy.

        Applies:
        - Grayscale conversion
        - Contrast enhancement
        - Noise reduction

        Args:
            image: PIL Image object

        Returns:
            PIL Image: Preprocessed image
        """
        try:
            from PIL import ImageEnhance, ImageFilter

            logger.debug("Preprocessing image for OCR")

            # Convert to grayscale (reduces complexity)
            if image.mode != 'L':
                image = image.convert('L')

            # Increase contrast (makes text stand out)
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)

            # Apply sharpening
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.5)

            # Denoise (median filter removes salt-and-pepper noise)
            image = image.filter(ImageFilter.MedianFilter(size=3))

            logger.debug("Image preprocessing completed")
            return image

        except Exception as e:
            logger.error(f"Image preprocessing failed: {e}", exc_info=True)
            # Return original image if preprocessing fails
            return image

    @staticmethod
    def preprocess_for_ocr_advanced(image):
        """
        Advanced preprocessing with adaptive thresholding.

        For documents with varying lighting conditions.

        Args:
            image: PIL Image object

        Returns:
            PIL Image: Preprocessed image with adaptive threshold
        """
        try:
            import cv2
            import numpy as np
            from PIL import Image

            logger.debug("Applying advanced OCR preprocessing")

            # Convert PIL Image to numpy array
            img_array = np.array(image)

            # Convert to grayscale if needed
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array

            # Apply adaptive thresholding
            binary = cv2.adaptiveThreshold(
                gray,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                11,
                2
            )

            # Convert back to PIL Image
            return Image.fromarray(binary)

        except ImportError:
            logger.warning("OpenCV not available, using basic preprocessing")
            return ImagePreprocessor.preprocess_for_ocr(image)
        except Exception as e:
            logger.error(f"Advanced preprocessing failed: {e}", exc_info=True)
            return ImagePreprocessor.preprocess_for_ocr(image)
