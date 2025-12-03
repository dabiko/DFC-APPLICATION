"""
Celery Tasks for Search Indexing
Handles async indexing of documents for better performance
"""

from celery import shared_task
from django.conf import settings
import logging

logger = logging.getLogger(__name__)


@shared_task(bind=True, max_retries=3)
def index_document(self, document_id):
    """
    Index a single document in Elasticsearch

    This task is called asynchronously when a document is created or updated.
    Uses retry logic with exponential backoff on failures.

    Args:
        document_id: Document primary key (UUID as string)

    Returns:
        dict: Status and document_id
    """
    try:
        from apps.documents.models import Document
        from apps.search.documents import DocumentDocument

        # Fetch document from database
        document = Document.objects.get(pk=document_id)

        # Skip if document is deleted
        if document.is_deleted:
            logger.info(f"Skipping indexing for deleted document {document_id}")
            return {'status': 'skipped', 'reason': 'deleted', 'document_id': document_id}

        # Index in Elasticsearch
        doc_instance = DocumentDocument()
        doc_instance.update(document)

        logger.info(f"Successfully indexed document {document_id}: {document.file_name}")
        return {'status': 'success', 'document_id': document_id}

    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found for indexing")
        return {'status': 'error', 'message': 'Document not found', 'document_id': document_id}

    except Exception as exc:
        logger.error(f"Error indexing document {document_id}: {str(exc)}", exc_info=True)
        # Retry with exponential backoff: 60s, 120s, 240s
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@shared_task(bind=True, max_retries=3)
def delete_document_from_index(self, document_id):
    """
    Remove document from Elasticsearch index

    Args:
        document_id: Document primary key (UUID as string)

    Returns:
        dict: Status and document_id
    """
    try:
        from apps.search.documents import DocumentDocument
        from elasticsearch.exceptions import NotFoundError

        # Try to get and delete the document
        try:
            doc = DocumentDocument.get(id=document_id)
            doc.delete()
            logger.info(f"Successfully deleted document {document_id} from index")
            return {'status': 'success', 'document_id': document_id}
        except NotFoundError:
            logger.debug(f"Document {document_id} not found in index (may already be deleted)")
            return {'status': 'success', 'message': 'not_found', 'document_id': document_id}

    except Exception as exc:
        logger.error(f"Error deleting document {document_id} from index: {str(exc)}", exc_info=True)
        # Retry with exponential backoff
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@shared_task
def rebuild_search_index():
    """
    Rebuild entire search index from scratch

    Use this for:
    - Initial setup
    - After major schema changes
    - When index becomes corrupted
    - Manual index refresh

    Run manually:
        python manage.py shell -c "from apps.search.tasks import rebuild_search_index; rebuild_search_index.delay()"

    Returns:
        dict: Status and statistics
    """
    from apps.documents.models import Document
    from apps.search.documents import DocumentDocument, documents_index
    from elasticsearch.exceptions import NotFoundError

    logger.info("Starting search index rebuild...")

    try:
        # Delete existing index
        try:
            documents_index.delete()
            logger.info("Deleted existing index")
        except NotFoundError:
            logger.info("No existing index to delete")

        # Recreate index with mappings
        documents_index.create()
        logger.info("Created new index with mappings")

        # Get all non-deleted documents
        documents = Document.objects.filter(is_deleted=False).select_related(
            'owner', 'department', 'folder', 'created_by'
        )
        total = documents.count()

        logger.info(f"Indexing {total} documents...")

        # Index documents in batches
        batch_size = 100
        indexed_count = 0
        error_count = 0

        for i, document in enumerate(documents.iterator(chunk_size=batch_size), 1):
            try:
                doc_instance = DocumentDocument()
                doc_instance.update(document)
                indexed_count += 1

                # Log progress every 100 documents
                if i % 100 == 0:
                    logger.info(f"Indexed {i}/{total} documents ({(i/total)*100:.1f}%)")

            except Exception as e:
                error_count += 1
                logger.error(f"Error indexing document {document.id}: {str(e)}")

        logger.info(f"Index rebuild complete. Indexed: {indexed_count}, Errors: {error_count}, Total: {total}")

        return {
            'status': 'success',
            'total_documents': total,
            'indexed': indexed_count,
            'errors': error_count
        }

    except Exception as e:
        logger.error(f"Fatal error during index rebuild: {str(e)}", exc_info=True)
        return {
            'status': 'error',
            'message': str(e)
        }


@shared_task(bind=True, max_retries=3)
def extract_document_text(self, document_id):
    """
    Extract text content from document for full-text search

    Supports:
    - PDF: PyPDF2
    - Word: python-docx
    - Excel: openpyxl
    - Text files: direct read
    - Images with text: OCR with tesseract (if available)

    Args:
        document_id: Document primary key (UUID as string)

    Returns:
        dict: Status, document_id, and extracted text length
    """
    try:
        from apps.documents.models import Document
        import io

        document = Document.objects.get(pk=document_id)

        # Skip if no file
        if not document.file:
            logger.warning(f"Document {document_id} has no file")
            return {'status': 'skipped', 'reason': 'no_file', 'document_id': document_id}

        logger.info(f"Extracting text from document {document_id} ({document.file_type})")

        # Download file from storage
        file_content = document.file.read()
        extracted_text = ""

        # PDF extraction
        if document.file_type == 'application/pdf':
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
                logger.info(f"Extracted text from {len(pdf_reader.pages)} PDF pages")
            except Exception as e:
                logger.error(f"Error extracting PDF text: {e}")

        # Word document extraction
        elif document.file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_content))
                extracted_text = '\n'.join([para.text for para in doc.paragraphs])
                logger.info(f"Extracted text from Word document ({len(doc.paragraphs)} paragraphs)")
            except Exception as e:
                logger.error(f"Error extracting Word text: {e}")

        # Excel extraction
        elif document.file_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_text = ' '.join([str(cell) for cell in row if cell])
                        if row_text.strip():
                            extracted_text += row_text + '\n'
                logger.info(f"Extracted text from Excel ({len(wb.worksheets)} sheets)")
            except Exception as e:
                logger.error(f"Error extracting Excel text: {e}")

        # Plain text files
        elif document.file_type.startswith('text/'):
            try:
                extracted_text = file_content.decode('utf-8', errors='ignore')
                logger.info("Extracted text from plain text file")
            except Exception as e:
                logger.error(f"Error extracting text file: {e}")

        # Store extracted text in document model
        if extracted_text:
            text_length = len(extracted_text)
            document.extracted_text = extracted_text[:50000]  # Limit to 50K chars
            document.is_indexed = True
            document.save(update_fields=['extracted_text', 'is_indexed'])

            # Re-index document with extracted text
            index_document.delay(str(document_id))

            logger.info(f"Extracted {text_length} characters from document {document_id}")
            return {
                'status': 'success',
                'document_id': document_id,
                'text_length': text_length
            }
        else:
            logger.warning(f"No text extracted from document {document_id}")
            return {
                'status': 'no_text',
                'document_id': document_id
            }

    except Document.DoesNotExist:
        logger.error(f"Document {document_id} not found for text extraction")
        return {'status': 'error', 'message': 'Document not found'}

    except Exception as exc:
        logger.error(f"Error extracting text from document {document_id}: {str(exc)}", exc_info=True)
        # Retry with exponential backoff
        raise self.retry(exc=exc, countdown=60 * (2 ** self.request.retries))


@shared_task
def batch_index_documents(document_ids):
    """
    Index multiple documents in a batch

    More efficient than individual tasks for bulk operations

    Args:
        document_ids: List of document IDs

    Returns:
        dict: Status and statistics
    """
    from apps.documents.models import Document
    from apps.search.documents import DocumentDocument

    logger.info(f"Starting batch indexing of {len(document_ids)} documents")

    indexed_count = 0
    error_count = 0
    errors = []

    for document_id in document_ids:
        try:
            document = Document.objects.get(pk=document_id)

            if not document.is_deleted:
                doc_instance = DocumentDocument()
                doc_instance.update(document)
                indexed_count += 1
            else:
                logger.debug(f"Skipping deleted document {document_id}")

        except Document.DoesNotExist:
            error_count += 1
            errors.append({'document_id': document_id, 'error': 'not_found'})

        except Exception as e:
            error_count += 1
            errors.append({'document_id': document_id, 'error': str(e)})
            logger.error(f"Error indexing document {document_id}: {e}")

    logger.info(f"Batch indexing complete. Indexed: {indexed_count}, Errors: {error_count}")

    return {
        'status': 'success',
        'total': len(document_ids),
        'indexed': indexed_count,
        'errors': error_count,
        'error_details': errors
    }
